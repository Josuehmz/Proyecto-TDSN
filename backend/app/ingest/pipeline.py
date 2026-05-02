from __future__ import annotations

import io
from dataclasses import dataclass
from uuid import UUID, uuid4

from sqlalchemy.orm import Session

from app.db.models import Chunk, Document
from app.logging_config import get_logger
from app.rag.chunker import chunk_text
from app.rag.embeddings import get_embeddings_model

logger = get_logger(__name__)


@dataclass(frozen=True)
class IngestedDocument:
    document_id: UUID
    title: str
    chunk_count: int


def _extract_text(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if lower.endswith(".docx"):
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    if lower.endswith((".html", ".htm")):
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "lxml")
        return soup.get_text("\n")
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:  # pragma: no cover
        return content.decode("latin-1", errors="ignore")


def ingest_document(
    db: Session,
    *,
    title: str,
    filename: str,
    content: bytes,
    mime_type: str,
    required_clearance: str,
    allowed_roles: list[str],
    allowed_departments: list[str],
) -> IngestedDocument:
    text = _extract_text(filename, content)
    if not text.strip():
        raise ValueError("El documento está vacío tras la extracción")

    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("El documento no generó chunks")

    doc = Document(
        id=uuid4(),
        title=title,
        source=filename,
        required_clearance=required_clearance,
        allowed_roles=allowed_roles,
        allowed_departments=allowed_departments,
        mime_type=mime_type,
        byte_size=len(content),
    )
    db.add(doc)
    db.flush()

    vectors = get_embeddings_model().embed([c.text for c in chunks])

    for c, vec in zip(chunks, vectors, strict=True):
        db.add(
            Chunk(
                document_id=doc.id,
                ordinal=c.ordinal,
                text=c.text,
                token_count=c.token_count,
                required_clearance=required_clearance,
                allowed_roles=allowed_roles,
                allowed_departments=allowed_departments,
                embedding=vec,
            )
        )

    logger.info(
        "ingest.done",
        document_id=str(doc.id),
        title=title,
        chunks=len(chunks),
        required_clearance=required_clearance,
    )
    return IngestedDocument(document_id=doc.id, title=title, chunk_count=len(chunks))
